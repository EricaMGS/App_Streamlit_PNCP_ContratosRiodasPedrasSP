import io
import datetime
import random
import re
import time

import pandas as pd
import requests
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from fpdf import FPDF


# ============================================================
# CONFIGURAÇÃO DO STREAMLIT
# ============================================================

st.set_page_config(
    page_title="Portal PNCP - Rio das Pedras/SP",
    page_icon="🏛️",
    layout="wide"
)


# ============================================================
# DADOS DO MUNICÍPIO
# ============================================================

CNPJ_RIO_DAS_PEDRAS = "44826840000183"
CODIGO_IBGE_RIO_DAS_PEDRAS = "3544004"
UF = "SP"

BASE_URL = "https://pncp.gov.br/api/consulta/v1"

# Limites de segurança
TAMANHO_PAGINA = 50
MAX_PAGINAS = 100
MAX_REGISTROS = 5000

# Timeout:
# conexão = 15 segundos
# leitura = 90 segundos
TIMEOUT = (15, 90)

# Número máximo de tentativas por requisição
MAX_TENTATIVAS = 5

# Pequena pausa entre páginas para não sobrecarregar a API
PAUSA_ENTRE_PAGINAS = 0.35


# ============================================================
# CABEÇALHOS HTTP
# ============================================================

HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/151.0 Safari/537.36"
    ),
    "Accept": "application/json",
    "Connection": "keep-alive",
}


# ============================================================
# FUNÇÕES UTILITÁRIAS
# ============================================================

def limpar_cnpj(valor):
    """Retorna somente os números de um CNPJ."""
    return re.sub(r"\D", "", str(valor or ""))


def formatar_moeda_br(valor):
    """
    Formata valores monetários no padrão brasileiro.

    Exemplos:
        15300296.03 -> R$ 15.300.296,03
        None -> R$ 0,00
        NaN -> R$ 0,00
    """
    try:
        if valor is None:
            return "R$ 0,00"

        if isinstance(valor, str):
            texto = valor.strip()

            if not texto:
                return "R$ 0,00"

            # Remove R$, espaços e separadores de milhar
            texto = texto.replace("R$", "").replace(" ", "")

            # Caso venha no formato brasileiro
            if "," in texto and "." in texto:
                texto = texto.replace(".", "").replace(",", ".")
            elif "," in texto:
                texto = texto.replace(",", ".")

            valor = float(texto)

        valor = float(valor)

        if pd.isna(valor):
            return "R$ 0,00"

        texto = f"{valor:,.2f}"

        return (
            f"R$ "
            f"{texto.replace(',', 'X').replace('.', ',').replace('X', '.')}"
        )

    except (ValueError, TypeError):
        return "R$ 0,00"


def converter_float(valor):
    """Converte diferentes formatos de número para float."""
    try:
        if valor is None:
            return 0.0

        if isinstance(valor, (int, float)):
            return float(valor)

        texto = str(valor).strip()

        if not texto:
            return 0.0

        texto = (
            texto
            .replace("R$", "")
            .replace(" ", "")
        )

        if "," in texto and "." in texto:
            texto = texto.replace(".", "").replace(",", ".")
        elif "," in texto:
            texto = texto.replace(",", ".")

        return float(texto)

    except (ValueError, TypeError):
        return 0.0


def formatar_data(valor):
    """Converte datas do PNCP para dd/mm/aaaa."""
    if valor is None:
        return ""

    texto = str(valor).strip()

    if not texto:
        return ""

    try:
        data = pd.to_datetime(texto, errors="coerce")

        if pd.isna(data):
            return texto

        return data.strftime("%d/%m/%Y")

    except Exception:
        return texto


def valor_campo(row, *campos, default="N/D"):
    """Retorna o primeiro campo existente e não vazio."""
    for campo in campos:
        valor = row.get(campo)

        if valor is not None:
            if isinstance(valor, float) and pd.isna(valor):
                continue

            if str(valor).strip() not in ("", "None", "nan"):
                return valor

    return default


def extrair_valor_dict(valor):
    """
    Converte estruturas aninhadas do PNCP para texto legível.
    """
    if not isinstance(valor, dict):
        return valor

    for chave in (
        "nome",
        "razaoSocial",
        "descricao",
        "descricaoNome",
        "municipioNome",
        "nomeOrgao",
        "nomeUnidade",
        "numero",
        "id",
    ):
        if chave in valor and valor[chave] not in (None, ""):
            return valor[chave]

    partes = []

    for chave, item in valor.items():
        if item not in (None, "") and not isinstance(item, (dict, list)):
            partes.append(f"{chave}: {item}")

    return " | ".join(partes) if partes else str(valor)


def tratar_dataframe(df):
    """
    Cria uma versão amigável do DataFrame sem alterar
    a estrutura original recebida da API.
    """
    if df.empty:
        return df.copy()

    resultado = df.copy()

    for coluna in resultado.columns:
        resultado[coluna] = resultado[coluna].apply(extrair_valor_dict)

    return resultado


def obter_id_pncp(row):
    """Identifica o número de controle PNCP."""
    return str(
        valor_campo(
            row,
            "numeroControlePNCP",
            "numeroControlePNCPAta",
            "numeroControlePNCPCompra",
            default="N/D",
        )
    )


def obter_processo(row):
    return str(
        valor_campo(
            row,
            "processo",
            "numeroProcesso",
            default="N/D",
        )
    )


def obter_objeto(row, tipo):
    if tipo == "Contratos":
        return str(
            valor_campo(
                row,
                "objetoContrato",
                "objeto",
                default="N/D",
            )
        )

    if tipo == "Atas de Registro de Preços":
        return str(
            valor_campo(
                row,
                "objetoContratacao",
                "objetoCompra",
                "objeto",
                default="N/D",
            )
        )

    return str(
        valor_campo(
            row,
            "objetoCompra",
            "objeto",
            "descricaoObjeto",
            default="N/D",
        )
    )


def obter_valor_registro(row, tipo):
    if tipo == "Contratos":
        return converter_float(
            valor_campo(
                row,
                "valorGlobal",
                "valorInicial",
                "valorContrato",
                default=0,
            )
        )

    return converter_float(
        valor_campo(
            row,
            "valorTotalHomologado",
            "valorTotalEstimado",
            "valorEstimado",
            default=0,
        )
    )


def obter_fornecedor(row):
    return str(
        valor_campo(
            row,
            "nomeRazaoSocialFornecedor",
            "fornecedor",
            "nomeFornecedor",
            default="N/D",
        )
    )


def obter_data_registro(row):
    return formatar_data(
        valor_campo(
            row,
            "dataPublicacaoPncp",
            "dataPublicacao",
            "dataAssinatura",
            "dataInclusao",
            "dataAtualizacao",
            default="",
        )
    )


# ============================================================
# SESSÃO HTTP
# ============================================================

@st.cache_resource
def criar_sessao_http():
    session = requests.Session()
    session.headers.update(HEADERS)
    return session


SESSION = criar_sessao_http()


# ============================================================
# EXTRATOR DE REGISTROS
# ============================================================

def extrair_registros(data):
    """
    Normaliza diferentes formatos de resposta do PNCP.
    """
    if data is None:
        return []

    if isinstance(data, list):
        return data

    if isinstance(data, dict):

        # Formatos mais comuns
        for chave in (
            "data",
            "items",
            "content",
            "dados",
            "results",
            "resultado",
        ):
            valor = data.get(chave)

            if isinstance(valor, list):
                return valor

        # Algumas respostas podem trazer estrutura de paginação
        # com "data" dentro de outro objeto.
        for valor in data.values():
            if isinstance(valor, list):
                return valor

    return []


# ============================================================
# ERRO PERSONALIZADO DA API
# ============================================================

class ErroPNCP(Exception):
    """Erro controlado de comunicação com o PNCP."""

    def __init__(self, mensagem, status_code=None, detalhes=None):
        super().__init__(mensagem)
        self.status_code = status_code
        self.detalhes = detalhes


# ============================================================
# CONSULTA INDIVIDUAL AO PNCP
# ============================================================

def consultar_pncp(url, params, max_tentativas=MAX_TENTATIVAS):
    """
    Executa uma requisição GET ao PNCP com:
    - timeout;
    - retry;
    - backoff exponencial;
    - jitter;
    - tratamento específico dos principais HTTPs.
    """

    ultimo_erro = None

    for tentativa in range(1, max_tentativas + 1):

        try:
            resposta = SESSION.get(
                url,
                params=params,
                timeout=TIMEOUT,
            )

            status = resposta.status_code

            # ------------------------------------------------
            # SUCESSO
            # ------------------------------------------------
            if status == 200:
                try:
                    return resposta.json()
                except ValueError as exc:
                    raise ErroPNCP(
                        "O PNCP respondeu, mas o conteúdo não é um JSON válido.",
                        status_code=status,
                        detalhes=resposta.text[:500],
                    ) from exc

            # ------------------------------------------------
            # SEM CONTEÚDO
            # ------------------------------------------------
            if status == 204:
                return []

            texto_resposta = resposta.text[:1000]

            # ------------------------------------------------
            # ERROS TEMPORÁRIOS
            # ------------------------------------------------
            if status in (429, 500, 502, 503, 504):

                ultimo_erro = ErroPNCP(
                    f"PNCP retornou HTTP {status}.",
                    status_code=status,
                    detalhes=texto_resposta,
                )

                if tentativa < max_tentativas:

                    espera_base = 2 ** (tentativa - 1)

                    # jitter entre 0 e 1,5 segundos
                    jitter = random.uniform(0, 1.5)

                    espera = min(
                        espera_base + jitter,
                        30
                    )

                    time.sleep(espera)
                    continue

                raise ultimo_erro

            # ------------------------------------------------
            # HTTP 400
            # ------------------------------------------------
            if status == 400:

                raise ErroPNCP(
                    (
                        "O PNCP rejeitou os parâmetros da consulta "
                        "(HTTP 400 - Bad Request)."
                    ),
                    status_code=status,
                    detalhes=texto_resposta,
                )

            # ------------------------------------------------
            # HTTP 422
            # ------------------------------------------------
            if status == 422:

                raise ErroPNCP(
                    (
                        "O PNCP considerou os parâmetros inválidos "
                        "(HTTP 422 - Unprocessable Entity)."
                    ),
                    status_code=status,
                    detalhes=texto_resposta,
                )

            # ------------------------------------------------
            # OUTROS ERROS
            # ------------------------------------------------
            raise ErroPNCP(
                f"O PNCP retornou HTTP {status}.",
                status_code=status,
                detalhes=texto_resposta,
            )

        except ErroPNCP:
            raise

        except (
            requests.exceptions.Timeout,
            requests.exceptions.ConnectionError,
            requests.exceptions.RequestException,
        ) as exc:

            ultimo_erro = exc

            if tentativa < max_tentativas:

                espera_base = 2 ** (tentativa - 1)
                jitter = random.uniform(0, 1.5)
                espera = min(espera_base + jitter, 30)

                time.sleep(espera)
                continue

            raise ErroPNCP(
                (
                    "Não foi possível estabelecer comunicação "
                    "com o PNCP após várias tentativas."
                ),
                detalhes=str(exc),
            ) from exc

    raise ErroPNCP(
        "Falha desconhecida na comunicação com o PNCP.",
        detalhes=str(ultimo_erro),
    )


# ============================================================
# PAGINAÇÃO SEQUENCIAL
# ============================================================

def consultar_paginas(
    url,
    params,
    max_paginas=MAX_PAGINAS,
    max_registros=MAX_REGISTROS,
):
    """
    Consulta as páginas sequencialmente.

    A consulta para quando:
    - não há registros;
    - a página vem incompleta;
    - o limite de páginas é atingido;
    - o limite de registros é atingido.
    """

    todos_registros = []

    estatisticas = {
        "paginas": 0,
        "registros": 0,
        "limite_atingido": False,
        "tempo": 0,
    }

    inicio = time.perf_counter()

    for pagina in range(1, max_paginas + 1):

        params_pagina = params.copy()
        params_pagina["pagina"] = pagina
        params_pagina["tamanhoPagina"] = TAMANHO_PAGINA

        dados = consultar_pncp(
            url,
            params_pagina,
        )

        registros = extrair_registros(dados)

        estatisticas["paginas"] = pagina

        if not registros:
            break

        todos_registros.extend(registros)

        estatisticas["registros"] = len(todos_registros)

        # Limite de segurança
        if len(todos_registros) >= max_registros:

            todos_registros = todos_registros[:max_registros]

            estatisticas["limite_atingido"] = True

            break

        # Se veio menos que o tamanho da página,
        # provavelmente chegamos à última página.
        if len(registros) < TAMANHO_PAGINA:
            break

        # Pequena pausa entre páginas
        if pagina < max_paginas:
            time.sleep(PAUSA_ENTRE_PAGINAS)

    estatisticas["tempo"] = time.perf_counter() - inicio

    return todos_registros, estatisticas


# ============================================================
# CACHE DAS CONSULTAS
# ============================================================

@st.cache_data(
    ttl=300,
    show_spinner=False,
)
def consultar_dados_cache(url, params_tuple):
    """
    Cache da consulta por 5 minutos.

    params_tuple é utilizado porque o Streamlit precisa
    de argumentos hashable.
    """

    params = dict(params_tuple)

    registros, estatisticas = consultar_paginas(
        url,
        params,
    )

    return registros, estatisticas


# ============================================================
# CONSULTA DE TERMOS / ADITIVOS DE CONTRATO
# ============================================================

def decompor_id_contrato(id_pncp):
    """
    Decompõe:
    CNPJ-2-SEQUENCIAL/ANO

    Exemplo:
    44826840000183-2-000001/2026
    """

    texto = str(id_pncp or "").strip()

    padrao = re.match(
        r"^(\d{14})-2-(\d+)/(\d{4})$",
        texto,
    )

    if not padrao:
        return None

    return {
        "cnpj": padrao.group(1),
        "sequencial": padrao.group(2),
        "ano": padrao.group(3),
    }


def consultar_detalhes_contrato(id_contrato):
    """
    Consulta os termos de um contrato utilizando
    a estrutura oficial CNPJ/ano/sequencial.
    """

    partes = decompor_id_contrato(id_contrato)

    if not partes:
        raise ErroPNCP(
            "Não foi possível interpretar o número de controle do contrato.",
            detalhes=str(id_contrato),
        )

    url = (
        f"https://pncp.gov.br/api/pncp/v1/orgaos/"
        f"{partes['cnpj']}/contratos/"
        f"{partes['ano']}/{partes['sequencial']}/termos"
    )

    try:
        dados = consultar_pncp(
            url,
            {},
            max_tentativas=3,
        )

        return extrair_registros(dados)

    except ErroPNCP:

        # Algumas instalações/versões podem não expor
        # os termos nesse endpoint. Retornamos erro controlado.
        raise


# ============================================================
# DADOS PARA RELATÓRIO
# ============================================================

def obter_dados_registro(row, tipo):

    id_pncp = obter_id_pncp(row)
    processo = obter_processo(row)
    objeto = obter_objeto(row, tipo)
    valor = obter_valor_registro(row, tipo)
    fornecedor = obter_fornecedor(row)
    data = obter_data_registro(row)

    if tipo == "Atas de Registro de Preços":

        vigencia_inicio = formatar_data(
            valor_campo(
                row,
                "vigenciaInicio",
                "dataVigenciaInicio",
                default="",
            )
        )

        vigencia_fim = formatar_data(
            valor_campo(
                row,
                "vigenciaFim",
                "dataVigenciaFim",
                default="",
            )
        )

        info_extra = (
            f"Vigência: {vigencia_inicio or 'N/D'} "
            f"até {vigencia_fim or 'N/D'}"
        )

    elif tipo == "Contratos":

        info_extra = (
            f"Fornecedor: {fornecedor} | "
            f"Valor: {formatar_moeda_br(valor)} | "
            f"Data: {data or 'N/D'}"
        )

    else:

        modalidade = valor_campo(
            row,
            "modalidadeNome",
            "modalidade",
            default="N/D",
        )

        info_extra = (
            f"Modalidade: {modalidade} | "
            f"Valor: {formatar_moeda_br(valor)} | "
            f"Data: {data or 'N/D'}"
        )

    return (
        id_pncp,
        processo,
        info_extra,
        objeto,
    )


# ============================================================
# LINK PARA O PNCP
# ============================================================

def gerar_link_pncp(row, tipo):

    id_pncp = obter_id_pncp(row)

    if id_pncp == "N/D":
        return ""

    # Contratação
    if tipo == "Editais e Avisos de Contratações":

        partes = id_pncp.split("-")

        if len(partes) >= 3:

            cnpj = partes[0]

            try:
                sequencial_ano = partes[-1]

                if "/" in sequencial_ano:
                    sequencial, ano = sequencial_ano.split("/")

                    return (
                        "https://pncp.gov.br/app/editais/"
                        f"{cnpj}/{ano}/{sequencial}"
                    )

            except Exception:
                pass

    # Contrato
    if tipo == "Contratos":

        partes = decompor_id_contrato(id_pncp)

        if partes:

            return (
                "https://pncp.gov.br/app/contratos/"
                f"{partes['cnpj']}/"
                f"{partes['ano']}/"
                f"{partes['sequencial']}"
            )

    # Ata
    if tipo == "Atas de Registro de Preços":

        partes = id_pncp.split("-")

        if len(partes) >= 3:

            cnpj = partes[0]

            try:
                sequencial_ano = partes[-1]

                if "/" in sequencial_ano:
                    sequencial, ano = sequencial_ano.split("/")

                    return (
                        "https://pncp.gov.br/app/atas/"
                        f"{cnpj}/{ano}/{sequencial}"
                    )

            except Exception:
                pass

    return ""


# ============================================================
# DATAFRAME RESUMIDO
# ============================================================

def criar_dataframe_resumo(df, tipo):

    if df.empty:
        return pd.DataFrame()

    linhas = []

    for _, row in df.iterrows():

        valor = obter_valor_registro(row, tipo)

        linha = {
            "ID PNCP": obter_id_pncp(row),
            "Processo": obter_processo(row),
            "Data": obter_data_registro(row),
            "Objeto": obter_objeto(row, tipo),
            "Fornecedor": obter_fornecedor(row),
            "Valor": valor,
            "Link PNCP": gerar_link_pncp(row, tipo),
        }

        linhas.append(linha)

    resumo = pd.DataFrame(linhas)

    if "Valor" in resumo.columns:
        resumo["Valor"] = pd.to_numeric(
            resumo["Valor"],
            errors="coerce",
        ).fillna(0)

    return resumo


# ============================================================
# CONFIGURAÇÃO DA PÁGINA
# ============================================================

st.title("🏛️ Contratações de Rio das Pedras/SP")

st.markdown(
    """
    Consulta integrada de **Contratos, Atas de Registro de Preços
    e Editais/Avisos de Contratações** diretamente no Portal Nacional
    de Contratações Públicas (PNCP).
    """
)

st.info(
    """
    ℹ️ **Atenção:** o PNCP é um serviço externo e pode apresentar
    lentidão, indisponibilidade ou respostas temporariamente vazias.
    Consultas com períodos maiores podem demorar mais. O sistema utiliza
    tentativas automáticas, paginação controlada e cache para reduzir
    os efeitos dessas instabilidades.
    """
)


# ============================================================
# SIDEBAR
# ============================================================

st.sidebar.header("⚙️ Parâmetros da Consulta")

tipo_consulta = st.sidebar.selectbox(
    "Selecione:",
    [
        "Contratos",
        "Atas de Registro de Preços",
        "Editais e Avisos de Contratações",
    ],
)


# ============================================================
# MODALIDADE
# ============================================================

modalidade_codigo = None

if tipo_consulta == "Editais e Avisos de Contratações":

    modalidade_opcoes = {
        "Pregão - Eletrônico (6)": 6,
        "Dispensa de Licitação (8)": 8,
        "Inexigibilidade (9)": 9,
        "Concorrência - Eletrônica (2)": 2,
    }

    mod_escolhida = st.sidebar.selectbox(
        "Modalidade:",
        list(modalidade_opcoes.keys()),
    )

    modalidade_codigo = modalidade_opcoes[mod_escolhida]


# ============================================================
# DATAS
# ============================================================

hoje = datetime.date.today()

data_inicio = st.sidebar.date_input(
    "Data Inicial",
    value=hoje - datetime.timedelta(days=90),
)

data_fim = st.sidebar.date_input(
    "Data Final",
    value=hoje,
)


# ============================================================
# VALIDAÇÃO DE DATAS
# ============================================================

if data_fim < data_inicio:

    st.sidebar.error(
        "⚠️ A Data Final não pode ser anterior à Data Inicial."
    )

    st.stop()


dias_periodo = (data_fim - data_inicio).days

if dias_periodo > 365:

    st.sidebar.error(
        "⚠️ O período não pode ser maior que 365 dias."
    )

    st.stop()


# ============================================================
# DIAGNÓSTICO
# ============================================================

diagnostico = st.sidebar.checkbox(
    "🔧 Mostrar diagnóstico técnico",
    value=False,
)


# ============================================================
# SESSION STATE
# ============================================================

if "df_resultado" not in st.session_state:
    st.session_state.df_resultado = None

if "estatisticas_consulta" not in st.session_state:
    st.session_state.estatisticas_consulta = None

if "parametros_consulta" not in st.session_state:
    st.session_state.parametros_consulta = None

if "tipo_anterior" not in st.session_state:
    st.session_state.tipo_anterior = tipo_consulta

if st.session_state.tipo_anterior != tipo_consulta:

    st.session_state.df_resultado = None
    st.session_state.estatisticas_consulta = None
    st.session_state.parametros_consulta = None

    st.session_state.tipo_anterior = tipo_consulta


# ============================================================
# ENDPOINTS
# ============================================================

endpoints = {
    "Contratos": f"{BASE_URL}/contratos",
    "Atas de Registro de Preços": f"{BASE_URL}/atas",
    "Editais e Avisos de Contratações": (
        f"{BASE_URL}/contratacoes/publicacao"
    ),
}


# ============================================================
# BOTÃO CONSULTA
# ============================================================

if st.sidebar.button(
    "🔎 Gerar Consulta",
    type="primary",
    use_container_width=True,
):

    endpoint = endpoints[tipo_consulta]

    # --------------------------------------------------------
    # Parâmetros comuns
    # --------------------------------------------------------

    params = {
        "dataInicial": data_inicio.strftime("%Y%m%d"),
        "dataFinal": data_fim.strftime("%Y%m%d"),
        "pagina": 1,
        "tamanhoPagina": TAMANHO_PAGINA,
    }

    # --------------------------------------------------------
    # Contratos
    # --------------------------------------------------------

    if tipo_consulta == "Contratos":

        params["cnpjOrgao"] = CNPJ_RIO_DAS_PEDRAS

    # --------------------------------------------------------
    # Atas
    # --------------------------------------------------------

    elif tipo_consulta == "Atas de Registro de Preços":

        params["cnpj"] = CNPJ_RIO_DAS_PEDRAS

    # --------------------------------------------------------
    # Editais
    # --------------------------------------------------------

    elif tipo_consulta == "Editais e Avisos de Contratações":

        params.update(
            {
                "codigoModalidadeContratacao": modalidade_codigo,
                "uf": UF,
                "codigoMunicipioIbge": (
                    CODIGO_IBGE_RIO_DAS_PEDRAS
                ),
                "cnpj": CNPJ_RIO_DAS_PEDRAS,
            }
        )

    # --------------------------------------------------------
    # Diagnóstico dos parâmetros
    # --------------------------------------------------------

    if diagnostico:

        with st.expander(
            "🔧 Parâmetros enviados ao PNCP",
            expanded=True,
        ):

            st.code(
                f"Endpoint:\n{endpoint}\n\n"
                f"Parâmetros:\n{params}",
                language="text",
            )

    try:

        inicio_consulta = time.perf_counter()

        with st.spinner(
            "🔄 Consultando o PNCP. "
            "A API pode apresentar lentidão..."
        ):

            # Tuple hashable para o cache
            params_tuple = tuple(
                sorted(params.items())
            )

            registros, estatisticas = consultar_dados_cache(
                endpoint,
                params_tuple,
            )

        tempo_total = time.perf_counter() - inicio_consulta

        df_temp = pd.DataFrame(registros)

        # ----------------------------------------------------
        # Tratamento
        # ----------------------------------------------------

        df_temp = tratar_dataframe(df_temp)

        # ----------------------------------------------------
        # Segurança adicional para Contratos
        # ----------------------------------------------------

        if (
            tipo_consulta == "Contratos"
            and not df_temp.empty
        ):

            possiveis_colunas = [
                "cnpjOrgao",
                "orgaoEntidade",
                "orgao",
                "unidadeOrgao",
            ]

            for coluna in possiveis_colunas:

                if coluna in df_temp.columns:

                    serie = (
                        df_temp[coluna]
                        .astype(str)
                        .str.replace(
                            r"\D",
                            "",
                            regex=True,
                        )
                    )

                    mask = serie.str.contains(
                        limpar_cnpj(CNPJ_RIO_DAS_PEDRAS),
                        na=False,
                    )

                    if mask.any():

                        df_temp = df_temp[mask].copy()
                        break

        # ----------------------------------------------------
        # Remove duplicidades
        # ----------------------------------------------------

        if not df_temp.empty:

            colunas_id = [
                coluna
                for coluna in (
                    "numeroControlePNCP",
                    "numeroControlePNCPAta",
                    "numeroControlePNCPCompra",
                )
                if coluna in df_temp.columns
            ]

            if colunas_id:

                coluna_id = colunas_id[0]

                df_temp = (
                    df_temp
                    .drop_duplicates(
                        subset=[coluna_id]
                    )
                    .reset_index(drop=True)
                )

        # ----------------------------------------------------
        # Salva estado
        # ----------------------------------------------------

        estatisticas["tempo_interface"] = tempo_total

        st.session_state.df_resultado = df_temp
        st.session_state.estatisticas_consulta = estatisticas
        st.session_state.parametros_consulta = params

    except ErroPNCP as erro:

        st.session_state.df_resultado = None
        st.session_state.estatisticas_consulta = None

        st.error(
            f"❌ {str(erro)}"
        )

        if erro.status_code == 400:

            st.warning(
                """
                ⚠️ O PNCP rejeitou os parâmetros da consulta.

                O sistema já utiliza **50 registros por página**,
                que é uma configuração conservadora para o endpoint
                de Editais/Avisos.

                Verifique principalmente o período e os filtros
                selecionados.
                """
            )

        elif erro.status_code in (
            429,
            500,
            502,
            503,
            504,
        ):

            st.warning(
                """
                🌐 O PNCP está apresentando uma resposta temporária
                de indisponibilidade ou sobrecarga.

                Tente novamente após alguns segundos ou utilize um
                período menor.
                """
            )

        if erro.detalhes and diagnostico:

            with st.expander(
                "🔍 Detalhes técnicos do erro"
            ):

                st.code(
                    str(erro.detalhes),
                    language="text",
                )

    except Exception as erro:

        st.session_state.df_resultado = None

        st.error(
            f"❌ Erro inesperado: {erro}"
        )

        if diagnostico:

            with st.expander(
                "🔍 Detalhes técnicos"
            ):

                st.exception(erro)


# ============================================================
# EXIBIÇÃO DOS RESULTADOS
# ============================================================

df = st.session_state.df_resultado
estatisticas = st.session_state.estatisticas_consulta


if df is not None and not df.empty:

    # ========================================================
    # CABEÇALHO DOS RESULTADOS
    # ========================================================

    st.success(
        f"📊 Consulta concluída: {len(df):,} registros encontrados."
        .replace(",", ".")
    )

    # ========================================================
    # ESTATÍSTICAS
    # ========================================================

    total_registros = len(df)

    coluna_valor = next(
        (
            coluna
            for coluna in (
                "valorGlobal",
                "valorInicial",
                "valorTotalHomologado",
                "valorTotalEstimado",
                "valorEstimado",
            )
            if coluna in df.columns
        ),
        None,
    )

    valor_total = 0.0

    if coluna_valor:

        valor_total = pd.to_numeric(
            df[coluna_valor],
            errors="coerce",
        ).fillna(0).sum()

    col_m1, col_m2, col_m3, col_m4 = st.columns(4)

    col_m1.metric(
        "📊 Registros",
        f"{total_registros:,}".replace(",", "."),
    )

    if coluna_valor:

        col_m2.metric(
            "💰 Valor Total",
            formatar_moeda_br(valor_total),
        )

    else:

        col_m2.metric(
            "💰 Valor Total",
            "N/D",
        )

    if estatisticas:

        col_m3.metric(
            "📄 Páginas",
            estatisticas.get("paginas", 0),
        )

        col_m4.metric(
            "⏱️ Tempo",
            f"{estatisticas.get('tempo', 0):.1f}s",
        )

    # ========================================================
    # AVISO DE LIMITE
    # ========================================================

    if (
        estatisticas
        and estatisticas.get("limite_atingido")
    ):

        st.warning(
            f"""
            ⚠️ A consulta atingiu o limite de segurança de
            {MAX_REGISTROS:,} registros.

            Reduza o período para garantir que todos os registros
            sejam carregados.
            """.replace(",", ".")
        )

    st.markdown("---")


    # ========================================================
    # DIAGNÓSTICO
    # ========================================================

    if diagnostico:

        with st.expander(
            "🔧 Diagnóstico da consulta",
            expanded=False,
        ):

            st.write(
                {
                    "Endpoint": endpoints[tipo_consulta],
                    "Tipo": tipo_consulta,
                    "Data inicial": data_inicio.strftime(
                        "%d/%m/%Y"
                    ),
                    "Data final": data_fim.strftime(
                        "%d/%m/%Y"
                    ),
                    "Registros": len(df),
                    "Páginas": (
                        estatisticas.get("paginas", 0)
                        if estatisticas
                        else 0
                    ),
                    "Tempo": (
                        f"{estatisticas.get('tempo', 0):.2f} s"
                        if estatisticas
                        else "N/D"
                    ),
                }
            )

            st.code(
                str(
                    st.session_state.parametros_consulta
                ),
                language="text",
            )


    # ========================================================
    # EXPORTAÇÃO
    # ========================================================

    st.markdown("### 📥 Opções de Exportação")

    cols = st.columns(4)

    nome = (
        tipo_consulta
        .replace(" ", "_")
        .replace("/", "_")
        .replace("ã", "a")
        .replace("ç", "c")
        .replace("õ", "o")
    )

    # --------------------------------------------------------
    # Excel
    # --------------------------------------------------------

    buffer_xlsx = io.BytesIO()

    df.to_excel(
        buffer_xlsx,
        index=False,
    )

    buffer_xlsx.seek(0)

    cols[0].download_button(
        "📊 Excel (.xlsx)",
        buffer_xlsx.getvalue(),
        f"{nome}_Rio_Das_Pedras.xlsx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "spreadsheetml.sheet"
        ),
        use_container_width=True,
    )

    # --------------------------------------------------------
    # CSV
    # --------------------------------------------------------

    csv_bytes = df.to_csv(
        index=False,
        encoding="utf-8-sig",
    ).encode("utf-8-sig")

    cols[1].download_button(
        "📄 CSV (.csv)",
        csv_bytes,
        f"{nome}_Rio_Das_Pedras.csv",
        mime="text/csv",
        use_container_width=True,
    )

    # --------------------------------------------------------
    # WORD
    # --------------------------------------------------------

    doc = Document()

    p_titulo = doc.add_paragraph()

    r_titulo = p_titulo.add_run(
        f"Relatório Executivo: {tipo_consulta}"
    )

    r_titulo.bold = True
    r_titulo.font.size = Pt(16)
    r_titulo.font.color.rgb = RGBColor(
        0,
        51,
        102,
    )

    doc.add_paragraph(
        "Município: Prefeitura Municipal de "
        "Rio das Pedras / SP"
    )

    doc.add_paragraph(
        f"Período: "
        f"{data_inicio.strftime('%d/%m/%Y')} "
        f"a "
        f"{data_fim.strftime('%d/%m/%Y')}"
    )

    doc.add_paragraph(
        f"Total de registros: {len(df)}"
    )

    if coluna_valor:

        doc.add_paragraph(
            f"Valor total envolvido: "
            f"{formatar_moeda_br(valor_total)}"
        )

    doc.add_paragraph(
        "Observação: o relatório Word apresenta "
        "os primeiros 50 registros. Excel e CSV "
        "contêm todos os registros carregados."
    )

    doc.add_heading(
        "Detalhamento dos Registros",
        level=2,
    )

    for posicao, (_, row) in enumerate(
        df.head(50).iterrows(),
        start=1,
    ):

        p_reg = doc.add_paragraph()

        p_reg.add_run(
            f"Item #{posicao}\n"
        ).bold = True

        id_pncp, processo, info_extra, objeto = (
            obter_dados_registro(
                row,
                tipo_consulta,
            )
        )

        p_reg.add_run(
            f"• ID Contratação PNCP: {id_pncp}\n"
        )

        p_reg.add_run(
            f"• Processo/Ref.: {processo}\n"
        )

        p_reg.add_run(
            f"• Detalhes: {info_extra}\n"
        )

        p_reg.add_run(
            f"• Objeto: {objeto}\n"
        )

        link = gerar_link_pncp(
            row,
            tipo_consulta,
        )

        if link:
            p_reg.add_run(
                f"• Link PNCP: {link}\n"
            )

        doc.add_paragraph(
            "-" * 50
        )

    buffer_docx = io.BytesIO()

    doc.save(buffer_docx)

    buffer_docx.seek(0)

    cols[2].download_button(
        "📝 Word (.docx)",
        buffer_docx.getvalue(),
        f"Relatorio_{nome}.docx",
        mime=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
        use_container_width=True,
    )

    # --------------------------------------------------------
    # PDF
    # --------------------------------------------------------

    pdf = FPDF()

    pdf.set_auto_page_break(
        auto=True,
        margin=15,
    )

    pdf.add_page()

    pdf.set_font(
        "Arial",
        "B",
        14,
    )

    pdf.cell(
        0,
        10,
        txt=f"Relatorio: {tipo_consulta}",
        ln=True,
        align="C",
    )

    pdf.set_font(
        "Arial",
        size=10,
    )

    pdf.cell(
        0,
        6,
        txt=(
            "Municipio: Prefeitura Municipal de "
            "Rio das Pedras / SP"
        ),
        ln=True,
    )

    pdf.cell(
        0,
        6,
        txt=(
            f"Periodo: "
            f"{data_inicio.strftime('%d/%m/%Y')} "
            f"a "
            f"{data_fim.strftime('%d/%m/%Y')}"
        ),
        ln=True,
    )

    pdf.cell(
        0,
        6,
        txt=f"Total: {len(df)} registros",
        ln=True,
    )

    if coluna_valor:

        valor_pdf = formatar_moeda_br(
            valor_total
        )

        pdf.cell(
            0,
            6,
            txt=f"Valor total: {valor_pdf}",
            ln=True,
        )

    pdf.ln(5)

    pdf.set_font(
        "Arial",
        "B",
        10,
    )

    pdf.cell(
        0,
        8,
        txt="Principais Registros:",
        ln=True,
    )

    pdf.set_font(
        "Arial",
        size=9,
    )

    for posicao, (_, row) in enumerate(
        df.head(30).iterrows(),
        start=1,
    ):

        id_pncp, processo, info_extra, objeto = (
            obter_dados_registro(
                row,
                tipo_consulta,
            )
        )

        bloco = (
            f"[{posicao}] "
            f"ID PNCP: {id_pncp} | "
            f"Proc.: {processo}\n"
            f"{info_extra}\n"
            f"Objeto: {objeto}"
        )

        bloco_limpo = (
            bloco
            .encode("latin-1", "replace")
            .decode("latin-1")
        )

        pdf.multi_cell(
            0,
            5,
            txt=bloco_limpo,
        )

        pdf.ln(3)

    pdf_bytes = pdf.output(
        dest="S"
    )

    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode(
            "latin-1"
        )

    cols[3].download_button(
        "📕 PDF (.pdf)",
        pdf_bytes,
        f"Relatorio_{nome}.pdf",
        mime="application/pdf",
        use_container_width=True,
    )

    st.caption(
        "Word: primeiros 50 registros | "
        "PDF: primeiros 30 registros | "
        "Excel/CSV: todos os registros carregados."
    )

    st.markdown("---")


    # ========================================================
    # FILTROS LOCAIS
    # ========================================================

    st.markdown("### 🔎 Filtros nos resultados")

    filtro_texto = st.text_input(
        "Pesquisar no objeto, processo ou fornecedor:",
        placeholder="Digite uma palavra ou expressão...",
    )

    df_filtrado = df.copy()

    if filtro_texto.strip():

        texto = filtro_texto.strip().lower()

        mascara = pd.Series(
            False,
            index=df_filtrado.index,
        )

        colunas_busca = [
            coluna
            for coluna in (
                "objetoCompra",
                "objetoContrato",
                "objeto",
                "processo",
                "numeroProcesso",
                "nomeRazaoSocialFornecedor",
                "nomeFornecedor",
            )
            if coluna in df_filtrado.columns
        ]

        for coluna in colunas_busca:

            mascara = (
                mascara
                | df_filtrado[coluna]
                .astype(str)
                .str.lower()
                .str.contains(
                    texto,
                    na=False,
                )
            )

        df_filtrado = df_filtrado[mascara]

    st.caption(
        f"Exibindo {len(df_filtrado)} de {len(df)} registros."
    )


    # ========================================================
    # ADITIVOS
    # ========================================================

    if tipo_consulta == "Contratos":

        st.markdown("---")

        st.markdown(
            "### 🔍 Consultar Aditivos / Termos do Contrato"
        )

        opcoes_contratos = []

        for _, row in df.iterrows():

            id_contrato = obter_id_pncp(row)

            processo = obter_processo(row)

            opcoes_contratos.append(
                f"{id_contrato} - Proc.: {processo}"
            )

        if opcoes_contratos:

            contrato_selecionado = st.selectbox(
                "Selecione um contrato:",
                opcoes_contratos,
            )

            numero_aditivo = st.text_input(
                "Número do termo/aditivo (opcional):",
                placeholder="Ex.: 01/2026",
            )

            if st.button(
                "🔎 Buscar Termos do Contrato",
                type="secondary",
            ):

                id_escolhido = contrato_selecionado.split(
                    " - "
                )[0]

                with st.spinner(
                    "🔄 Consultando termos do contrato no PNCP..."
                ):

                    try:

                        termos = consultar_detalhes_contrato(
                            id_escolhido
                        )

                        if termos:

                            # Filtro opcional
                            if numero_aditivo.strip():

                                termo_busca = (
                                    numero_aditivo
                                    .strip()
                                    .lower()
                                )

                                termos = [
                                    termo
                                    for termo in termos
                                    if termo_busca
                                    in str(termo).lower()
                                ]

                            if termos:

                                st.success(
                                    f"Encontrados {len(termos)} "
                                    "termo(s)/documento(s)."
                                )

                                tabela_termos = []

                                for termo in termos:

                                    tabela_termos.append(
                                        {
                                            "Número": valor_campo(
                                                termo,
                                                "numero",
                                                "numeroTermo",
                                                default="S/N",
                                            ),
                                            "Tipo": valor_campo(
                                                termo,
                                                "tipoDocumentoNome",
                                                "tipoTermo",
                                                "tipo",
                                                default="N/D",
                                            ),
                                            "Data": formatar_data(
                                                valor_campo(
                                                    termo,
                                                    "dataPublicacaoPncp",
                                                    "dataPublicacao",
                                                    "dataAssinatura",
                                                    default="",
                                                )
                                            ),
                                            "Objeto": valor_campo(
                                                termo,
                                                "objeto",
                                                "descricao",
                                                default="N/D",
                                            ),
                                            "Valor Acrescido": (
                                                formatar_moeda_br(
                                                    converter_float(
                                                        valor_campo(
                                                            termo,
                                                            "valorAcrescido",
                                                            default=0,
                                                        )
                                                    )
                                                )
                                            ),
                                        }
                                    )

                                st.dataframe(
                                    pd.DataFrame(
                                        tabela_termos
                                    ),
                                    use_container_width=True,
                                    hide_index=True,
                                )

                            else:

                                st.warning(
                                    "Nenhum termo encontrado "
                                    "com o número informado."
                                )

                        else:

                            st.info(
                                "ℹ️ Nenhum termo/aditivo foi "
                                "retornado para este contrato."
                            )

                    except ErroPNCP as erro:

                        st.error(
                            f"❌ Não foi possível consultar os "
                            f"termos: {erro}"
                        )

                        if diagnostico and erro.detalhes:

                            st.code(
                                str(erro.detalhes),
                                language="text",
                            )

        else:

            st.info(
                "ℹ️ Nenhum contrato disponível para consulta."
            )


    # ========================================================
    # ANÁLISE GRÁFICA
    # ========================================================

    st.markdown("---")

    st.markdown(
        "### 📈 Análise dos Registros"
    )

    coluna_data = next(
        (
            coluna
            for coluna in (
                "dataPublicacaoPncp",
                "dataPublicacao",
                "dataAssinatura",
                "dataInclusao",
                "dataAtualizacao",
            )
            if coluna in df.columns
        ),
        None,
    )

    if coluna_data:

        try:

            df_grafico = df.copy()

            df_grafico["_data"] = pd.to_datetime(
                df_grafico[coluna_data],
                errors="coerce",
            )

            df_grafico = df_grafico[
                df_grafico["_data"].notna()
            ].copy()

            if not df_grafico.empty:

                df_grafico["_mes_ano"] = (
                    df_grafico["_data"]
                    .dt.to_period("M")
                    .astype(str)
                )

                contagem_mes = (
                    df_grafico["_mes_ano"]
                    .value_counts()
                    .sort_index()
                )

                st.markdown(
                    "#### Quantidade de registros por mês"
                )

                st.bar_chart(
                    contagem_mes
                )

                # --------------------------------------------
                # Gráfico de valores
                # --------------------------------------------

                if coluna_valor:

                    df_grafico["_valor"] = pd.to_numeric(
                        df_grafico[coluna_valor],
                        errors="coerce",
                    ).fillna(0)

                    valor_mes = (
                        df_grafico
                        .groupby("_mes_ano")["_valor"]
                        .sum()
                        .sort_index()
                    )

                    st.markdown(
                        "#### Valor total por mês"
                    )

                    st.bar_chart(
                        valor_mes
                    )

            else:

                st.info(
                    "ℹ️ Não há datas válidas para gerar "
                    "o gráfico temporal."
                )

        except Exception:

            st.info(
                "ℹ️ Não foi possível gerar automaticamente "
                "o gráfico temporal."
            )

    else:

        st.info(
            "ℹ️ Coluna de data não encontrada."
        )


    # ========================================================
    # TABELA RESUMIDA
    # ========================================================

    st.markdown("---")

    st.markdown(
        "### 📋 Tabela Resumida"
    )

    df_resumo = criar_dataframe_resumo(
        df_filtrado,
        tipo_consulta,
    )

    if not df_resumo.empty:

        # Formata valor somente para apresentação
        df_resumo_exibicao = df_resumo.copy()

        if "Valor" in df_resumo_exibicao.columns:

            df_resumo_exibicao["Valor"] = (
                df_resumo_exibicao["Valor"]
                .apply(formatar_moeda_br)
            )

        st.dataframe(
            df_resumo_exibicao,
            use_container_width=True,
            hide_index=True,
            column_config={
                "Link PNCP": st.column_config.LinkColumn(
                    "🔗 PNCP",
                    display_text="Abrir no PNCP",
                ),
                "Objeto": st.column_config.TextColumn(
                    "Objeto",
                    width="large",
                ),
            },
        )

    else:

        st.info(
            "ℹ️ Nenhum registro corresponde ao filtro."
        )


    # ========================================================
    # DADOS COMPLETOS
    # ========================================================

    with st.expander(
        "🧾 Ver dados completos retornados pelo PNCP"
    ):

        st.dataframe(
            df_filtrado,
            use_container_width=True,
            hide_index=True,
        )


# ============================================================
# NENHUM RESULTADO
# ============================================================

elif (
    st.session_state.df_resultado is not None
    and st.session_state.df_resultado.empty
):

    st.warning(
        """
        ⚠️ Nenhum registro encontrado para
        Rio das Pedras/SP no período selecionado.

        Isso não significa necessariamente que não existam
        contratações. O PNCP pode retornar uma consulta vazia
        durante períodos de instabilidade ou quando o filtro
        selecionado não possui registros publicados no período.
        """
    )


# ============================================================
# RODAPÉ
# ============================================================

st.markdown("---")

st.caption(
    "Portal PNCP - Rio das Pedras/SP | "
    "Consulta de dados públicos do PNCP"
)
